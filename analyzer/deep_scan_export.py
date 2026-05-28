"""Export the deep-scan result as Markdown, DOCX, or PDF.

All three renderers consume the same dict shape returned by
``deep_scan.analyze_deep`` / ``deep_scan.run_from_text``.
"""

from __future__ import annotations

from datetime import datetime
from io import BytesIO, StringIO


# ============================================================
# Markdown
# ============================================================

def render_md(result: dict, topic: str | None = None) -> str:
    out = StringIO()
    out.write("# Deep Scan Report\n\n")
    out.write(f"*Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}*\n\n")
    if topic:
        out.write(f"**Topic hint:** {topic}\n\n")

    out.write("## Summary\n\n")
    out.write(f"- Word count: {result['n_words']}\n")
    out.write(f"- Metrics evaluated: {result['metric_count']}\n")
    out.write(
        f"- Total markers: **{result['total_markers_raw']}** "
        f"(**{result['total_markers_per_500']} per 500 words**)\n"
    )
    out.write(f"- Density classification: **{result['density'].upper()}**\n\n")

    out.write(
        f"Source for all markers: [{result['wikipedia_url']}]"
        f"({result['wikipedia_url']})\n\n"
    )

    out.write(
        "_This report describes stylistic patterns. It does **not** assert "
        "authorship, AI generation, or academic integrity conclusions. "
        "Humans use every one of these patterns; high clusters across "
        "multiple metrics are the diagnostically meaningful signal._\n\n"
    )

    out.write("## Per-Metric Findings\n\n")
    for key, m in result["metrics"].items():
        out.write(f"### {m['name']}\n\n")
        if m.get("description"):
            out.write(f"*{m['description']}*\n\n")
        out.write(
            f"- **Count:** {m['raw_count']} | **Rate:** {m['per_500']} per 500 words\n"
        )
        out.write(f"- **Tier:** {m.get('tier', '?')} | **Key:** `{key}`\n")
        out.write(f"- **Wikipedia:** [↗ section]({m['wikipedia_section']})\n")
        if m.get("notes"):
            out.write(f"- **Notes:** {m['notes']}\n")
        out.write("\n")

        if m.get("top_hits"):
            out.write("**Top hits:**\n\n")
            for hit, count in m["top_hits"][:10]:
                if isinstance(hit, str):
                    out.write(f"- `{hit}` ({count}×)\n")
                else:
                    out.write(f"- `{hit}`: {count}\n")
            out.write("\n")

        examples = m.get("examples") or []
        if examples:
            out.write("**Examples:**\n\n")
            for ex in examples[:5]:
                if isinstance(ex, dict):
                    if "sentence" in ex:
                        out.write(f"- \"{ex['sentence']}\"\n")
                    elif "items" in ex:
                        out.write(f"- triplet: `{' / '.join(ex['items'])}`\n")
                    elif "term" in ex:
                        out.write(f"- `{ex['term']}`\n")
                    elif "members_present" in ex:
                        out.write(
                            f"- paragraph {ex['paragraph']}, cluster "
                            f"*{ex['cluster']}*: "
                            f"{', '.join(f'`{m}`' for m in ex['members_present'])}\n"
                        )
                    elif "from_paragraph" in ex:
                        out.write(
                            f"- paragraphs {ex['from_paragraph']} → "
                            f"{ex['to_paragraph']}: "
                            f"{ex['dimensions_shifted']} dimensions shifted\n"
                        )
                    else:
                        out.write(f"- {ex}\n")
                else:
                    out.write(f"- {ex}\n")
            out.write("\n")

        details = m.get("details") or {}
        if details:
            out.write("**Details:**\n\n")
            for k, v in details.items():
                out.write(f"- {k}: `{v}`\n")
            out.write("\n")

    return out.getvalue()


# ============================================================
# DOCX
# ============================================================

def render_docx(result: dict, topic: str | None = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Title
    title = doc.add_heading("Deep Scan Report", level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").italic = True

    if topic:
        p = doc.add_paragraph()
        p.add_run("Topic hint: ").bold = True
        p.add_run(topic)

    # Summary
    doc.add_heading("Summary", level=1)
    doc.add_paragraph(f"Word count: {result['n_words']}", style="List Bullet")
    doc.add_paragraph(f"Metrics evaluated: {result['metric_count']}", style="List Bullet")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(f"Total markers: ").bold = False
    r = p.add_run(f"{result['total_markers_raw']}")
    r.bold = True
    p.add_run(f" ({result['total_markers_per_500']} per 500 words)")
    p = doc.add_paragraph(style="List Bullet")
    p.add_run("Density classification: ").bold = False
    r = p.add_run(result["density"].upper())
    r.bold = True

    doc.add_paragraph(
        "This report describes stylistic patterns. It does NOT assert "
        "authorship, AI generation, or academic integrity conclusions. "
        "Humans use every one of these patterns; high clusters across "
        "multiple metrics are the diagnostically meaningful signal."
    ).italic = True

    # Per-metric
    doc.add_heading("Per-Metric Findings", level=1)

    for key, m in result["metrics"].items():
        doc.add_heading(m["name"], level=2)

        if m.get("description"):
            d = doc.add_paragraph(m["description"])
            d.runs[0].italic = True

        # Headline numbers
        p = doc.add_paragraph()
        p.add_run("Count: ").bold = True
        p.add_run(f"{m['raw_count']}   ")
        p.add_run("Rate: ").bold = True
        p.add_run(f"{m['per_500']} per 500 words   ")
        p.add_run("Tier: ").bold = True
        p.add_run(f"{m.get('tier', '?')}")

        p = doc.add_paragraph()
        p.add_run("Wikipedia section: ").bold = True
        p.add_run(m["wikipedia_section"])

        if m.get("notes"):
            p = doc.add_paragraph()
            p.add_run("Notes: ").bold = True
            p.add_run(m["notes"])

        if m.get("top_hits"):
            doc.add_paragraph("Top hits:").runs[0].bold = True
            for hit, count in m["top_hits"][:10]:
                if isinstance(hit, str):
                    doc.add_paragraph(f"{hit} ({count}×)", style="List Bullet")
                else:
                    doc.add_paragraph(f"{hit}: {count}", style="List Bullet")

        examples = m.get("examples") or []
        if examples:
            doc.add_paragraph("Examples:").runs[0].bold = True
            for ex in examples[:5]:
                if isinstance(ex, dict):
                    if "sentence" in ex:
                        doc.add_paragraph(f"\"{ex['sentence']}\"", style="List Bullet")
                    elif "items" in ex:
                        doc.add_paragraph(
                            f"triplet: {' / '.join(ex['items'])}",
                            style="List Bullet",
                        )
                    elif "term" in ex:
                        doc.add_paragraph(ex["term"], style="List Bullet")
                    elif "members_present" in ex:
                        doc.add_paragraph(
                            f"paragraph {ex['paragraph']}, cluster "
                            f"{ex['cluster']}: {', '.join(ex['members_present'])}",
                            style="List Bullet",
                        )
                    elif "from_paragraph" in ex:
                        doc.add_paragraph(
                            f"paragraphs {ex['from_paragraph']} → "
                            f"{ex['to_paragraph']}: "
                            f"{ex['dimensions_shifted']} dimensions shifted",
                            style="List Bullet",
                        )

        details = m.get("details") or {}
        if details:
            doc.add_paragraph("Details:").runs[0].bold = True
            for k, v in details.items():
                doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# PDF (reportlab)
# ============================================================

def render_pdf(result: dict, topic: str | None = None) -> bytes:
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import (
        ListFlowable, ListItem, PageBreak, Paragraph, SimpleDocTemplate,
        Spacer,
    )

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=letter,
        leftMargin=0.85 * inch, rightMargin=0.85 * inch,
        topMargin=0.85 * inch, bottomMargin=0.85 * inch,
        title="Deep Scan Report",
    )
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle(
            "title", parent=base["Title"], fontName="Helvetica-Bold",
            fontSize=22, spaceAfter=6,
        ),
        "meta": ParagraphStyle(
            "meta", parent=base["Normal"], fontName="Helvetica-Oblique",
            fontSize=9, textColor="#666666", spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "h1", parent=base["Heading1"], fontName="Helvetica-Bold",
            fontSize=15, spaceBefore=18, spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "h2", parent=base["Heading2"], fontName="Helvetica-Bold",
            fontSize=12, spaceBefore=12, spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "body", parent=base["Normal"], fontName="Helvetica",
            fontSize=10, leading=14, spaceAfter=6,
        ),
        "bullet": ParagraphStyle(
            "bullet", parent=base["Normal"], fontName="Helvetica",
            fontSize=10, leading=13, leftIndent=14, bulletIndent=0,
        ),
        "italic": ParagraphStyle(
            "italic", parent=base["Normal"], fontName="Helvetica-Oblique",
            fontSize=10, leading=14, textColor="#444444", spaceAfter=6,
        ),
        "disclaimer": ParagraphStyle(
            "disclaimer", parent=base["Normal"], fontName="Helvetica-Oblique",
            fontSize=9, leading=12, textColor="#555555",
            spaceBefore=10, spaceAfter=14,
        ),
        "url": ParagraphStyle(
            "url", parent=base["Normal"], fontName="Helvetica",
            fontSize=8, textColor="#555555", leading=11, spaceAfter=4,
        ),
    }

    def esc(s: str) -> str:
        return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    story = []
    story.append(Paragraph("Deep Scan Report", styles["title"]))
    story.append(Paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
        styles["meta"],
    ))

    if topic:
        story.append(Paragraph(f"<b>Topic hint:</b> {esc(topic)}", styles["body"]))

    # Summary
    story.append(Paragraph("Summary", styles["h1"]))
    summary_items = [
        f"Word count: {result['n_words']}",
        f"Metrics evaluated: {result['metric_count']}",
        f"Total markers: <b>{result['total_markers_raw']}</b> "
        f"({result['total_markers_per_500']} per 500 words)",
        f"Density classification: <b>{result['density'].upper()}</b>",
    ]
    story.append(ListFlowable(
        [ListItem(Paragraph(s, styles["bullet"])) for s in summary_items],
        bulletType="bullet", leftIndent=18,
    ))
    story.append(Spacer(1, 6))
    story.append(Paragraph(
        "This report describes stylistic patterns. It does NOT assert "
        "authorship, AI generation, or academic integrity conclusions. "
        "Humans use every one of these patterns; high clusters across "
        "multiple metrics are the diagnostically meaningful signal.",
        styles["disclaimer"],
    ))
    story.append(Paragraph(
        f"Source: <a href='{result['wikipedia_url']}'>{result['wikipedia_url']}</a>",
        styles["url"],
    ))

    # Per-metric
    story.append(PageBreak())
    story.append(Paragraph("Per-Metric Findings", styles["h1"]))

    for key, m in result["metrics"].items():
        story.append(Paragraph(esc(m["name"]), styles["h2"]))

        if m.get("description"):
            story.append(Paragraph(esc(m["description"]), styles["italic"]))

        head = (
            f"<b>Count:</b> {m['raw_count']} &nbsp; "
            f"<b>Rate:</b> {m['per_500']} per 500w &nbsp; "
            f"<b>Tier:</b> {esc(str(m.get('tier', '?')))}"
        )
        story.append(Paragraph(head, styles["body"]))

        story.append(Paragraph(
            f"<a href='{m['wikipedia_section']}'>↗ Wikipedia section</a>",
            styles["url"],
        ))

        if m.get("notes"):
            story.append(Paragraph(f"<i>{esc(m['notes'])}</i>", styles["body"]))

        if m.get("top_hits"):
            hits_text = "<b>Top hits:</b> " + ", ".join(
                f"{esc(str(h))} ({c})" if isinstance(h, str) else f"{esc(str(h))}: {c}"
                for h, c in m["top_hits"][:8]
            )
            story.append(Paragraph(hits_text, styles["body"]))

        examples = m.get("examples") or []
        if examples:
            items: list[ListItem] = []
            for ex in examples[:3]:
                if isinstance(ex, dict):
                    if "sentence" in ex:
                        items.append(ListItem(Paragraph(
                            f"\"{esc(str(ex['sentence']))}\"",
                            styles["bullet"],
                        )))
                    elif "items" in ex:
                        items.append(ListItem(Paragraph(
                            f"triplet: {esc(' / '.join(ex['items']))}",
                            styles["bullet"],
                        )))
                    elif "term" in ex:
                        items.append(ListItem(Paragraph(
                            esc(ex["term"]), styles["bullet"],
                        )))
                    elif "members_present" in ex:
                        items.append(ListItem(Paragraph(
                            f"paragraph {ex['paragraph']}, cluster "
                            f"<i>{esc(ex['cluster'])}</i>: "
                            f"{esc(', '.join(ex['members_present']))}",
                            styles["bullet"],
                        )))
                    elif "from_paragraph" in ex:
                        items.append(ListItem(Paragraph(
                            f"paragraphs {ex['from_paragraph']} → "
                            f"{ex['to_paragraph']}: "
                            f"{ex['dimensions_shifted']} dimensions shifted",
                            styles["bullet"],
                        )))
            if items:
                story.append(Paragraph("<b>Examples:</b>", styles["body"]))
                story.append(ListFlowable(
                    items, bulletType="bullet", leftIndent=18,
                ))

        story.append(Spacer(1, 6))

    doc.build(story)
    return buf.getvalue()
