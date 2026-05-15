"""Extract plain text from uploaded files.

Supported formats: .txt, .md, .docx, .pdf, .json, .html / .htm.

The extracted text is fed into the stylometric pipeline exactly as if it had
been pasted into the textarea. We do not preserve formatting beyond paragraph
breaks, since the analyzer works on plain prose.
"""

from __future__ import annotations

import io
import json
import os
import re
from html.parser import HTMLParser


SUPPORTED_EXTENSIONS = {"txt", "md", "markdown", "docx", "pdf", "json", "html", "htm"}


class UnsupportedFileError(ValueError):
    """Raised when the upload's extension is not supported."""


class ExtractionError(RuntimeError):
    """Raised when a known format fails to parse."""


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = (os.path.splitext(filename)[1] or "").lower().lstrip(".")
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileError(
            f"Unsupported file type: .{ext or '(none)'}. "
            f"Supported: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
        )

    if ext in ("txt", "md", "markdown"):
        return _extract_plain(file_bytes)
    if ext == "docx":
        return _extract_docx(file_bytes)
    if ext == "pdf":
        return _extract_pdf(file_bytes)
    if ext == "json":
        return _extract_json(file_bytes)
    if ext in ("html", "htm"):
        return _extract_html(file_bytes)
    raise UnsupportedFileError(f"Unsupported file type: .{ext}")


def _extract_plain(file_bytes: bytes) -> str:
    # Try UTF-8, then fall back to latin-1 (which never errors).
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


def _extract_docx(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("python-docx is not installed.") from exc
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ExtractionError(f"Could not open .docx file: {exc}") from exc

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also pull table cell text so footnoted material isn't silently dropped.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    text = "\n\n".join(paragraphs).strip()
    if not text:
        raise ExtractionError(
            ".docx contained no extractable text "
            "(it may consist entirely of images or be empty)."
        )
    return text


def _extract_pdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise ExtractionError("pypdf is not installed.") from exc
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ExtractionError(f"Could not open PDF: {exc}") from exc
    if reader.is_encrypted:
        # pypdf may be able to open an empty-password PDF; try once.
        try:
            reader.decrypt("")
        except Exception:
            raise ExtractionError("PDF is encrypted; remove password before uploading.")

    pages: list[str] = []
    for page in reader.pages:
        try:
            page_text = page.extract_text() or ""
        except Exception:
            page_text = ""
        if page_text.strip():
            pages.append(page_text.strip())

    text = "\n\n".join(pages).strip()
    if not text:
        raise ExtractionError(
            "PDF contained no extractable text "
            "(it may be a scanned image; OCR is not supported)."
        )
    return text


def _extract_json(file_bytes: bytes) -> str:
    try:
        data = json.loads(file_bytes.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ExtractionError(f"Could not parse JSON: {exc}") from exc

    strings: list[str] = []

    def _walk(node) -> None:
        if isinstance(node, str):
            s = node.strip()
            if s:
                strings.append(s)
        elif isinstance(node, dict):
            for v in node.values():
                _walk(v)
        elif isinstance(node, list):
            for item in node:
                _walk(item)

    _walk(data)
    text = "\n\n".join(strings).strip()
    if not text:
        raise ExtractionError("JSON contained no string values.")
    return text


class _HTMLTextExtractor(HTMLParser):
    """Strip tags, drop <script>/<style> content, treat block tags as paragraphs."""

    _BLOCK_TAGS = {
        "p", "div", "section", "article", "header", "footer", "main",
        "h1", "h2", "h3", "h4", "h5", "h6",
        "li", "blockquote", "pre", "br", "hr", "tr",
    }
    _SKIP_TAGS = {"script", "style", "noscript", "iframe"}

    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._skip_depth = 0

    def handle_starttag(self, tag: str, attrs) -> None:  # noqa: D401
        if tag in self._SKIP_TAGS:
            self._skip_depth += 1
        if tag in self._BLOCK_TAGS:
            self._chunks.append("\n\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in self._SKIP_TAGS and self._skip_depth > 0:
            self._skip_depth -= 1
        if tag in self._BLOCK_TAGS:
            self._chunks.append("\n\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth == 0:
            self._chunks.append(data)

    @property
    def text(self) -> str:
        joined = "".join(self._chunks)
        # Collapse whitespace within lines, preserve blank-line separators.
        joined = re.sub(r"[ \t]+", " ", joined)
        joined = re.sub(r"\n{3,}", "\n\n", joined)
        return joined.strip()


def _extract_html(file_bytes: bytes) -> str:
    raw = _extract_plain(file_bytes)
    parser = _HTMLTextExtractor()
    parser.feed(raw)
    parser.close()
    text = parser.text
    if not text:
        raise ExtractionError("HTML contained no extractable text.")
    return text
